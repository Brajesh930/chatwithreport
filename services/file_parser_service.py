from pathlib import Path
from datetime import datetime
from helpers.config import Config
from helpers.logger import Logger
from helpers.file_helper import FileHelper

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

class FileParserService:
    """Service for parsing uploaded files"""

    def __init__(self):
        Config.load()
        storage_path = Path(Config.get('STORAGE_PATH', './storage'))
        self.upload_dir = storage_path / 'uploads'
        self.parsed_dir = storage_path / 'parsed'
        self.parsed_dir.mkdir(parents=True, exist_ok=True)

    def parse(self, uploaded_file_path, original_filename):
        """Parse uploaded file and save as normalized text"""
        uploaded_file_path = Path(uploaded_file_path)
        
        if not uploaded_file_path.exists():
            return {'success': False, 'error': 'Uploaded file not found'}

        extension = uploaded_file_path.suffix.lower().lstrip('.')
        
        # Extract text based on file type
        text = None
        parser_used = ''
        extraction_note = ''

        try:
            if extension == 'txt':
                text = self._parse_txt(uploaded_file_path)
                parser_used = 'Plain Text Parser'
            elif extension == 'md':
                text = self._parse_markdown(uploaded_file_path)
                parser_used = 'Markdown Parser'
            elif extension == 'pdf':
                result = self._parse_pdf(uploaded_file_path)
                text = result['text']
                parser_used = 'PDF Parser'
                extraction_note = result.get('note', '')
            elif extension == 'docx':
                result = self._parse_docx(uploaded_file_path)
                text = result['text']
                parser_used = 'DOCX Parser'
                extraction_note = result.get('note', '')
            else:
                return {'success': False, 'error': f'Unsupported file format: {extension}'}
        except Exception as e:
            Logger.error(f'Parse error for file: {original_filename}', {'error': str(e)})
            return {'success': False, 'error': f'Failed to parse file: {str(e)}'}

        if not text:
            return {'success': False, 'error': 'No text could be extracted from file'}

        # Create normalized text file
        file_size = uploaded_file_path.stat().st_size
        upload_time = datetime.fromtimestamp(uploaded_file_path.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')

        normalized_content = self._build_normalized_format(
            original_filename,
            uploaded_file_path.name,
            extension,
            upload_time,
            file_size,
            parser_used,
            text,
            extraction_note
        )

        # Save normalized file
        parsed_filename = f"{uploaded_file_path.stem}_parsed.txt"
        parsed_file_path = self.parsed_dir / parsed_filename

        try:
            with open(parsed_file_path, 'w', encoding='utf-8') as f:
                f.write(normalized_content)

            Logger.info('File parsed successfully', {
                'original': original_filename,
                'parsed': parsed_filename,
                'parser': parser_used
            })

            return {
                'success': True,
                'original_filename': original_filename,
                'parsed_filename': parsed_filename,
                'parser_used': parser_used,
                'parsed_path': str(parsed_file_path)
            }
        except Exception as e:
            Logger.error('Failed to save parsed file', {'error': str(e)})
            return {'success': False, 'error': 'Failed to save parsed file'}

    def _parse_txt(self, file_path):
        """Parse plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _parse_markdown(self, file_path):
        """Parse markdown file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _parse_pdf(self, file_path):
        """Parse PDF file"""
        if PdfReader is None:
            return {
                'text': '',
                'note': 'PyPDF2 not installed - PDF parsing limited'
            }

        try:
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

            text = '\n\n'.join(text_parts)
            return {
                'text': text if text else 'No text found in PDF',
                'note': f'Extracted from {len(reader.pages)} pages'
            }
        except Exception as e:
            return {
                'text': '',
                'note': f'PDF parsing error: {str(e)}'
            }

    def _parse_docx(self, file_path):
        """Parse DOCX file"""
        if Document is None:
            return {
                'text': '',
                'note': 'python-docx not installed - DOCX parsing not available'
            }

        try:
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            text = '\n'.join(text_parts)
            return {
                'text': text if text else 'No text found in DOCX',
                'note': f'Extracted paragraphs and tables'
            }
        except Exception as e:
            return {
                'text': '',
                'note': f'DOCX parsing error: {str(e)}'
            }

    def _build_normalized_format(self, original_filename, stored_filename, extension, upload_time, 
                                  file_size, parser_used, text, extraction_note):
        """Build normalized text format with metadata"""
        header = f"""================================================================================
DOCUMENT CHAT SYSTEM - NORMALIZED TEXT FILE
================================================================================
Original Filename: {original_filename}
Stored Filename: {stored_filename}
File Type: {extension.upper()}
Upload Time: {upload_time}
File Size: {FileHelper.format_file_size(file_size)}
Parser Used: {parser_used}
{f'Extraction Note: {extraction_note}' if extraction_note else ''}
================================================================================

"""
        return header + text
